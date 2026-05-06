# -*- coding: utf-8 -*-
"""把所有 docx 模板转成 HTML 写入数据库"""
import subprocess, os
from docx import Document

MYSQL = r'D:\phpstudy_pro\Extensions\MySQL8.0.12\bin\mysql.exe'

# 查询所有有文件但没有html_content的模板
result = subprocess.run(
    [MYSQL, '-u', 'root', '-proot', '-N', '-e',
     "SELECT id, file_path FROM trade_platform.biz_template WHERE file_path IS NOT NULL AND (html_content IS NULL OR html_content='') AND deleted=0;"],
    capture_output=True, text=True, encoding='utf-8'
)

for line in result.stdout.strip().split('\n'):
    if not line.strip():
        continue
    parts = line.split('\t')
    if len(parts) < 2:
        continue
    tid, fpath = parts[0].strip(), parts[1].strip()
    fpath = fpath.replace('\\', '/')

    if not os.path.exists(fpath):
        print(f"  ID={tid}: file not found: {fpath}")
        continue

    if not fpath.endswith('.docx'):
        print(f"  ID={tid}: not docx, skip")
        continue

    print(f"  ID={tid}: converting {os.path.basename(fpath)}...")
    doc = Document(fpath)
    html = ['<div style="font-family:SimSun,serif;line-height:1.8">']
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            html.append('<p><br></p>')
            continue
        style = para.style.name if para.style else ''
        if 'Title' in style or 'Heading 1' in style:
            html.append(f'<h1>{text}</h1>')
        elif 'Heading 2' in style:
            html.append(f'<h2>{text}</h2>')
        else:
            parts_html = ''
            for run in para.runs:
                t = run.text
                if not t: continue
                if run.bold: t = f'<strong>{t}</strong>'
                if run.italic: t = f'<em>{t}</em>'
                if run.underline: t = f'<u>{t}</u>'
                parts_html += t
            if not parts_html: parts_html = text
            html.append(f'<p>{parts_html}</p>')

    for table in doc.tables:
        html.append('<table style="width:100%;border-collapse:collapse;margin:10px 0">')
        for i, row in enumerate(table.rows):
            html.append('<tr>')
            for cell in row.cells:
                tag = 'th' if i == 0 else 'td'
                html.append(f'<{tag} style="border:1px solid #333;padding:5px 8px">{cell.text.strip()}</{tag}>')
            html.append('</tr>')
        html.append('</table>')
    html.append('</div>')

    content = '\n'.join(html)
    escaped = content.replace('\\', '\\\\').replace("'", "\\'")
    sql = f"UPDATE biz_template SET html_content='{escaped}' WHERE id={tid};"

    sql_file = f'D:/work/product/autoWorkPlatform-v1.0.0/tests/temp_update_{tid}.sql'
    with open(sql_file, 'w', encoding='utf-8') as f:
        f.write(sql)

    r = subprocess.run(
        [MYSQL, '-u', 'root', '-proot', 'trade_platform', '-e', f'source {sql_file}'],
        capture_output=True, text=True, encoding='utf-8'
    )
    os.remove(sql_file)
    if r.returncode == 0:
        print(f"  ID={tid}: OK ({len(content)} chars)")
    else:
        print(f"  ID={tid}: FAIL: {r.stderr[:100]}")

print("Done!")
